from docx import Document
from datetime import datetime

date = datetime.now()
date_formated = date.strftime("%d/%m/%Y")
archive = "references/testplan.docx"

class DocWord:
    def __init__(self, test_plan):
        self.document = Document(archive)
        self.test_plan = test_plan
        self.test_plan_data = {}
        self.test_cases_count = 1

    def set_placeholders(self):
        self.test_plan_data = {
            "[TEAM_NAME]": f'{self.test_plan.team_name}',
            "[PROJECT_NAME]": f'{self.test_plan.project_name}',
            "[FUNCTIONALITY_NAME]": f'{self.test_plan.functionality_name}',
            "[GENTI_HISTORY]": f'{self.test_plan.genti_history}',
            "[DATE]": f'{date_formated}'
        }

    def write_placeholders(self):
        for paragraph in self.document.paragraphs:
            for placeholder, valor in self.test_plan_data.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, valor)

    def write_title(self, test_case_title):
        self.document.add_paragraph(f"{self.test_cases_count} - {test_case_title}", style='List Paragraph')
        self.test_cases_count += 1

    def write_steps(self, test_case_steps):
        for step in test_case_steps:
            self.document.add_paragraph(step, style='List Paragraph')

    def write_testcase(self, test_case):
        self.write_title(test_case.title)
        self.write_steps(test_case.steps)

    def create_testcases(self, test_cases):
        for test_case in test_cases.test_cases_list:
            self.write_testcase(test_case)

    def create_document(self):
        self.document.save("casos_de_teste_gerados.docx")
        print("Documento gerado com sucesso!")